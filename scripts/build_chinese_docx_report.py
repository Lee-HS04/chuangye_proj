from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SRC = Path(
    r"C:\Users\GanJX\Documents\WeChat Files\wxid_7c7d1hfn0wjt12\FileStorage\File\2026-06\New Microsoft Word Document (2).docx"
)
OUT = Path(r"C:\Users\GanJX\Desktop\chuangye_proj\多智能体强化学习项目研究路线详细报告.docx")

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)
BLACK = RGBColor(0x00, 0x00, 0x00)
FILL = "F2F4F7"
CALLOUT = "F4F6F9"


def set_east_asian_font(run, latin="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_run(run, size=None, bold=None, color=None, latin="Calibri", east_asia="Microsoft YaHei"):
    set_east_asian_font(run, latin, east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_paragraph(doc, text="", style=None, bold_lead=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run(r1, size=11, bold=True, color=BLACK)
        r2 = p.add_run(text[len(bold_lead) :])
        set_run(r2, size=11, color=BLACK)
    else:
        r = p.add_run(text)
        set_run(r, size=11, color=BLACK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run(r, size={1: 16, 2: 13, 3: 12}[level], bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run(r, size=11, color=BLACK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run(r, size=11, color=BLACK)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        set_cell_shading(hdr[i], FILL)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run(r, size=10.5, bold=True, color=BLACK)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == len(row) - 1 or len(text) > 16 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_run(r, size=10, color=BLACK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    r = p.add_run(title + "：")
    set_run(r, size=10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(body)
    set_run(r, size=10.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("研究路线报告 | 多智能体强化学习与机器人协同控制")
    set_run(r, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("生成日期：2026-06-22")
    set_run(r, size=9, color=MUTED)


def extract_source():
    source = Document(str(SRC))
    paras = [p.text.strip() for p in source.paragraphs if p.text.strip()]
    tables = []
    for table in source.tables:
        tables.append([[cell.text.strip() for cell in row.cells] for row in table.rows])
    with ZipFile(SRC) as zf:
        media_count = len([n for n in zf.namelist() if n.startswith("word/media/")])
        comments = [n for n in zf.namelist() if "comments" in n.lower()]
    props = source.core_properties
    return {
        "paragraphs": paras,
        "tables": tables,
        "paragraph_count": len(source.paragraphs),
        "nonempty_count": len(paras),
        "table_count": len(tables),
        "media_count": media_count,
        "comments_count": len(comments),
        "author": props.author or "未填写",
        "created": props.created,
        "modified": props.modified,
        "revision": props.revision,
    }


def main():
    data = extract_source()
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("多智能体强化学习项目研究路线详细报告")
    set_run(r, size=23, bold=True, color=BLACK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("基于源 Word 文档的内容梳理、技术分析与执行建议")
    set_run(r, size=14, color=MUTED)

    for label, value in [
        ("源文件", SRC.name),
        ("报告语言", "中文"),
        ("文档主题", "MARL 研究路线、仿真复现、机器人系统迁移与硬件部署"),
        ("生成日期", "2026-06-22"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label + "：")
        set_run(r, size=11, bold=True, color=BLACK)
        r = p.add_run(value)
        set_run(r, size=11, color=BLACK)

    add_callout(
        doc,
        "核心判断",
        "原文档的主线非常明确：不要急于把 GitHub 代码直接接入机械狗或小车，而是先用 PettingZoo/MPE、MAPPO、MADDPG、QMIX 等代码完成研究验证、仿真复现和算法理解，再迁移到“机械狗 + 无人机 + 机械车 + 鲲鹏平台 + ESP32”的分层协同控制系统。",
    )

    add_heading(doc, "一、文档基本信息", 1)
    rows = [
        ("段落数量", f"{data['paragraph_count']} 个段落，其中 {data['nonempty_count']} 个为非空段落"),
        ("表格数量", f"{data['table_count']} 个表格，主要用于算法优先级、文献关系、仿真映射和 RL 概念映射"),
        ("图片/媒体", f"{data['media_count']} 个，源文档未包含图片类素材"),
        ("批注/评论", f"{data['comments_count']} 个结构化批注部件，源文档未检测到批注"),
        ("作者信息", data["author"]),
        ("创建/修改时间", f"{data['created']} / {data['modified']}"),
        ("修订号", str(data["revision"])),
    ]
    add_table(doc, ["项目", "说明"], rows, [2200, 7160])

    add_heading(doc, "二、原文档主题概述", 1)
    add_paragraph(
        doc,
        "这份 Word 文档是一份面向机器人协同控制课题的研究路线建议稿。它围绕“多智能体强化学习如何真正服务硬件项目”展开，重点不是立即进行硬件部署，而是先通过仿真环境理解多智能体系统中的 observation、action、reward、policy 与 env，再逐步迁移到实际系统。",
    )
    add_paragraph(
        doc,
        "文档中反复强调的系统架构可以概括为：机械狗、OpenMV 与树莓派负责边缘感知；鲲鹏平台承担中心决策和策略推理；ESP32/PWM 负责底层执行控制；最终目标是完成目标跟随、多节点通信、协同控制和系统集成验证。",
    )

    add_heading(doc, "三、技术路线详细分析", 1)
    add_heading(doc, "1. 代码路线选择", 2)
    add_paragraph(
        doc,
        "原文建议将 PettingZoo/MPE 放在第一优先级，而不是一开始就运行 QMIX + SMAC。这个判断合理，因为 SMAC/StarCraft 环境安装较重，且任务形态与真实机器人硬件之间的迁移距离较远。PettingZoo/MPE 更轻量，适合作为机械狗、无人机和机械车协同环境的原型接口。",
    )
    add_table(
        doc,
        ["优先级", "代码/算法", "报告解读"],
        [
            ("第一", "PettingZoo / MPE", "作为环境接口和轻量仿真入口，最适合先理解多智能体交互流程。"),
            ("第二", "MAPPO", "作为稳定 baseline，用于验证合作控制策略和训练流程。"),
            ("第三", "MADDPG", "适合连续动作控制，例如速度、转角、舵机角度等输出。"),
            ("第四", "QMIX / PyMARL", "更适合作为理论对照、相关工作和合作型任务分配补充。"),
        ],
        [1300, 2500, 5560],
    )

    add_heading(doc, "2. 老师要求“先做 research”时应提交的成果", 2)
    for item in [
        "文献表：整理 MADDPG、QMIX、MAPPO、PettingZoo/MPE 的算法思想、适用场景以及与本项目的对应关系。",
        "代码复现实验：先跑 simple_spread 或 cooperative navigation 这类轻量环境，观察训练曲线、奖励变化和 agent 行为。",
        "项目迁移方案：说明如何从仿真训练过渡到鲲鹏平台高层决策，再由 ESP32 执行 PWM 底层控制。",
    ]:
        add_bullet(doc, item)
    add_paragraph(
        doc,
        "这三类交付物的价值在于，它们可以向导师证明项目不是停留在硬件堆叠，而是在用可复现的算法实验支撑后续系统集成。",
    )

    add_heading(doc, "3. GitHub 代码的正确使用方式", 2)
    for step in [
        "下载代码并完成依赖安装。",
        "先运行官方 demo，确认环境、训练脚本和输出结果可用。",
        "重点阅读 env.reset()、env.step(actions) 和 reward 定义。",
        "将官方环境的 observation/action/reward 映射为机器人项目中的状态、动作和奖励。",
        "训练模型后，把策略部署到鲲鹏平台，由鲲鹏输出速度、转角或任务指令给 ESP32。",
    ]:
        add_number(doc, step)
    add_callout(
        doc,
        "实践重点",
        "强化学习代码不需要一开始全部看懂，应先抓住 reset、step、reward 和 policy.learn 这条主线。只要理解一次状态进入策略、动作进入环境、奖励反馈更新模型的闭环，就能逐步改成自己的机器人场景。",
    )

    add_heading(doc, "四、算法与项目关系矩阵", 1)
    add_table(
        doc,
        ["算法/环境", "方法特点", "解决问题", "与项目关系"],
        [
            ("MADDPG", "集中式 critic + 分布式 actor", "多智能体连续控制", "对应小车速度、舵机角度、目标跟随等连续动作。"),
            ("QMIX", "价值分解", "合作型任务分配", "对应多设备协同调度，可作为理论增强和对比实验。"),
            ("MAPPO", "多智能体 PPO", "稳定、强 baseline", "适合作为仿真训练主方法。"),
            ("PettingZoo/MPE", "标准多智能体环境接口", "提供轻量训练环境", "可改造成机械狗、UAV 和小车仿真系统。"),
        ],
        [1600, 2200, 2200, 3360],
    )

    add_heading(doc, "五、系统架构迁移方案", 1)
    add_paragraph(
        doc,
        "原文提出的工程路线可以整理为分层智能体协同框架。上层负责任务协同和调度，中层负责连续运动控制，底层负责真实硬件执行。这种分层方式有利于把强化学习策略和硬件控制解耦，降低直接部署带来的风险。",
    )
    add_table(
        doc,
        ["层级", "建议方法/组件", "主要作用", "输出"],
        [
            ("任务协同层", "QMIX 或 MAPPO", "决定机械狗、无人机、小车在任务中的角色与调度关系", "任务状态、目标分配、协同策略"),
            ("运动控制层", "MADDPG 或 MAPPO", "处理连续动作控制与目标跟随", "速度 v、转角 theta、动作 command"),
            ("工程执行层", "ESP32 + PWM", "执行舵机、电机和底层控制指令", "PWM 信号与真实机械动作"),
            ("感知输入层", "OpenMV、AprilTag、树莓派", "识别目标位置并生成环境状态", "目标坐标、距离、角度误差"),
        ],
        [1700, 2200, 3560, 1900],
    )

    add_heading(doc, "六、仿真环境设计建议", 1)
    add_table(
        doc,
        ["仿真元素", "项目映射", "设计说明"],
        [
            ("agent_0", "机械狗", "可承担巡逻、跟随、停止或接近目标任务。"),
            ("agent_1", "无人机", "可承担搜索、悬停、标记或辅助感知任务。"),
            ("agent_2", "机械车", "可承担移动执行、目标跟随和地面转运任务。"),
            ("target/landmark", "AprilTag 或工业目标点", "作为目标定位、跟随和任务完成判定依据。"),
            ("reward", "距离、碰撞、抖动、任务完成", "建议使用距离误差惩罚、控制平滑惩罚、碰撞惩罚和完成奖励组合。"),
        ],
        [1700, 2600, 5060],
    )
    add_paragraph(
        doc,
        "建议的 observation 可以从最小可用状态开始，例如 [x_self, y_self, x_target, y_target, distance, angle_error]。动作空间可以先离散化，便于快速验证；后续再扩展到连续速度、转角和姿态控制。",
    )

    add_heading(doc, "七、四周研究执行计划", 1)
    add_table(
        doc,
        ["周期", "主要任务", "输出成果", "验收标准"],
        [
            ("第 1 周", "文献调研：MADDPG、QMIX、MAPPO、PettingZoo/MPE", "文献对比表与项目适配分析", "能说明每种方法适合解决什么问题。"),
            ("第 2 周", "跑通 PettingZoo simple_spread_v3 或 MAPPO MPE demo", "训练曲线、reward 变化、agent 运动截图、实验说明", "官方环境可复现，能解释 observation/action/reward。"),
            ("第 3 周", "改造成机械狗/UAV/小车简化仿真环境", "自定义环境原型与奖励函数", "能完成 reset、step、reward 的闭环。"),
            ("第 4 周", "写入学推或中期报告", "研究进展、算法路线、系统迁移计划", "能清楚表达从规则控制到学习型协同决策的升级路径。"),
        ],
        [1200, 3160, 2700, 2300],
    )

    add_heading(doc, "八、风险与改进建议", 1)
    add_table(
        doc,
        ["风险点", "影响", "建议处理方式"],
        [
            ("一开始直接上硬件", "调试成本高，难以判断问题来自算法、通信还是控制", "先完成仿真闭环，再逐步接入真实传感与执行层。"),
            ("直接跑 QMIX + SMAC", "环境重、迁移弱，容易耗费时间", "把 QMIX 放在理论对比或后期增强，不作为第一步。"),
            ("奖励函数过于复杂", "训练不稳定，结果难解释", "先用距离误差、碰撞惩罚、完成奖励构成最小奖励函数。"),
            ("算法和硬件接口耦合过紧", "后续替换算法或硬件困难", "保持鲲鹏平台输出高层 action，ESP32 只负责底层执行。"),
        ],
        [2200, 3000, 4160],
    )

    add_heading(doc, "九、可直接用于汇报的结论", 1)
    add_paragraph(
        doc,
        "本项目已明确以多智能体强化学习作为协同控制策略优化方向。当前阶段应优先完成 PettingZoo/MPE 仿真环境复现，理解多智能体系统中的状态、动作和奖励建模方式；随后以 MAPPO 作为主要 baseline，以 MADDPG 支持连续控制，以 QMIX 作为合作型任务分配和理论对比补充。",
    )
    add_paragraph(
        doc,
        "在工程落地上，鲲鹏平台不应直接替代 ESP32 的底层控制，而应作为高层智能决策模块，输出目标速度、转向角或任务指令；ESP32 继续执行 PWM 控制。这样可以形成“边缘感知—中心决策—底层执行”的清晰系统架构，也与原项目申请中的技术目标保持一致。",
    )
    add_callout(
        doc,
        "一句话总结",
        "GitHub 代码不是直接拿来控制机器人，而是先用来复现实验、理解 MARL 的 observation/action/reward，再把这个结构迁移成“机械狗 + UAV + 小车 + 鲲鹏平台”的协同控制场景。",
    )

    add_heading(doc, "附录：源文档表格内容整理", 1)
    for idx, table in enumerate(data["tables"], 1):
        if not table:
            continue
        add_heading(doc, f"附表 {idx}", 2)
        headers = table[0]
        rows = table[1:]
        widths = [9360 // len(headers)] * len(headers)
        widths[-1] += 9360 - sum(widths)
        add_table(doc, headers, rows, widths)

    doc.core_properties.title = "多智能体强化学习项目研究路线详细报告"
    doc.core_properties.subject = "Word 源文档中文详细报告"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "多智能体强化学习, MAPPO, MADDPG, QMIX, PettingZoo, MPE, 鲲鹏平台, ESP32"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(str(OUT).encode("unicode_escape").decode("ascii"))


if __name__ == "__main__":
    main()
